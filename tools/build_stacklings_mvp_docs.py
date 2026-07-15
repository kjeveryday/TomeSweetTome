from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C2CC"


def set_run_font(run, *, size=None, bold=None, italic=None, color=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches):
    widths_dxa = [round(width * 1440) for width in widths_inches]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2B579A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.extend([r_fonts, color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\[[^\]]+\]\(https?://[^)]+\)|`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text, *, size=11, bold=False, color=None):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, bold=bold, color=color)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(9, size - 0.5), bold=bold, color=color, name="Courier New")
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, bold=bold, color=color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def table_widths(headers):
    n = len(headers)
    lowered = [h.lower() for h in headers]
    if n == 4 and "required for core loop" in lowered:
        return [1.45, 0.85, 2.0, 2.2]
    if n == 4:
        return [1.35, 1.15, 2.0, 2.0]
    if n == 3 and "source" in lowered:
        return [1.45, 3.35, 1.7]
    if n == 3:
        return [1.6, 2.8, 2.1]
    if n == 2:
        return [1.7, 4.8]
    return [6.5 / n] * n


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    set_table_borders(table)
    set_table_geometry(table, table_widths(headers))
    for row_index, source_row in enumerate(rows):
        target_row = table.rows[row_index]
        tr_pr = target_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, value in enumerate(source_row):
            cell = target_row.cells[col_index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            clear_paragraph(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, value, size=9.2, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, LIGHT_GRAY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167
    nested_bullet = styles["List Bullet 2"]
    nested_bullet.font.name = FONT
    nested_bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    nested_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    nested_bullet.font.size = Pt(11)
    nested_bullet.paragraph_format.left_indent = Inches(0.75)
    nested_bullet.paragraph_format.first_line_indent = Inches(-0.25)
    nested_bullet.paragraph_format.space_after = Pt(4)
    nested_bullet.paragraph_format.line_spacing = 1.167


def set_page(doc, short_label):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(short_label.upper())
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def render_markdown(md_path, output_path, short_label):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    set_page(doc, short_label)
    in_code = False
    code_lines = []
    index = 0
    first_heading = True
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.15)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.0
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F6F8FA")
                p_pr.append(shd)
                run = p.add_run("\n".join(code_lines))
                set_run_font(run, size=8.8, name="Courier New")
                in_code = False
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                if not re.match(r"^\|?\s*:?-+", lines[index]):
                    rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            if rows:
                add_table(doc, rows)
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if first_heading:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                set_run_font(run, size=23, bold=True, color=RGBColor(0, 0, 0))
                first_heading = False
            else:
                doc.add_heading(text, level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif bullet_match := re.match(r"^(\s*)- ", line):
            nested = bool(bullet_match.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.75 if nested else 0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4 if nested else 6)
            p.paragraph_format.line_spacing = 1.167
            p.paragraph_format.keep_together = True
            bullet_run = p.add_run("•\u00A0")
            set_run_font(bullet_run, size=11)
            add_inline(p, line[bullet_match.end():].strip())
        elif number_match := re.match(r"^(\d+)\. ", line):
            # Render the source number directly so each Markdown list restarts exactly
            # where the author intended instead of inheriting Word's prior list state.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.167
            p.paragraph_format.keep_together = True
            number_run = p.add_run(f"{number_match.group(1)}. ")
            set_run_font(number_run, size=11)
            add_inline(p, line[number_match.end():].strip())
        else:
            p = doc.add_paragraph()
            add_inline(p, line.strip())
        index += 1
    doc.save(output_path)


def main():
    pairs = [
        (ROOT / "Stacklings MVP PRD v0.3.md", ROOT / "Stacklings MVP PRD v0.3.docx", "Stacklings MVP Product Requirements"),
        (ROOT / "Stacklings MVP Parking Lot v0.1.md", ROOT / "Stacklings MVP Parking Lot v0.1.docx", "Stacklings MVP Parking Lot"),
    ]
    for source, output, label in pairs:
        render_markdown(source, output, label)
        print(output)


if __name__ == "__main__":
    main()
